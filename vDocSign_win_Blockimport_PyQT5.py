import sys
import os
import glob
import subprocess

try:
    from PyQt5 import QtWidgets, QtGui
except ModuleNotFoundError:
    print("PyQt5 is not installed. Please install it using 'pip install PyQt5'")
    sys.exit(1)
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocxProcessorApp(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Docx Stempel und Signatur")
        self.setGeometry(100, 100, 600, 400)

        # Initialisiere Pfade
        self.input_folder = ""
        self.output_folder = ""
        self.image_stamp_path = ""

        # Menüband erstellen
        self.create_menu()

        # Zentrales Widget
        self.central_widget = QtWidgets.QWidget()
        self.setCentralWidget(self.central_widget)

        # Layout
        layout = QtWidgets.QVBoxLayout()
        self.central_widget.setLayout(layout)

        # Labels und Buttons für Verzeichnisse
        self.input_label = QtWidgets.QLabel("Eingabeordner: Nicht ausgewählt")
        layout.addWidget(self.input_label)

        self.select_input_button = QtWidgets.QPushButton("Eingabeordner auswählen")
        self.select_input_button.clicked.connect(self.select_input_folder)
        layout.addWidget(self.select_input_button)

        self.output_label = QtWidgets.QLabel("Ausgabeordner: Nicht ausgewählt")
        layout.addWidget(self.output_label)

        self.select_output_button = QtWidgets.QPushButton("Ausgabeordner auswählen")
        self.select_output_button.clicked.connect(self.select_output_folder)
        layout.addWidget(self.select_output_button)

        self.image_label = QtWidgets.QLabel("Stempelbild: Nicht ausgewählt")
        layout.addWidget(self.image_label)

        self.select_image_button = QtWidgets.QPushButton("Stempelbild auswählen")
        self.select_image_button.clicked.connect(self.select_image)
        layout.addWidget(self.select_image_button)

        # Button zum Starten der Verarbeitung
        self.process_button = QtWidgets.QPushButton("DOCX-Dateien verarbeiten")
        self.process_button.clicked.connect(self.process_files)
        layout.addWidget(self.process_button)

        # Statusanzeige
        self.status_bar = QtWidgets.QStatusBar()
        self.setStatusBar(self.status_bar)

    def create_menu(self):
        menubar = self.menuBar()
        file_menu = menubar.addMenu("&Datei")

        exit_action = QtWidgets.QAction("Beenden", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

    def select_input_folder(self):
        folder = QtWidgets.QFileDialog.getExistingDirectory(self, "Eingabeordner auswählen")
        if folder:
            self.input_folder = folder
            self.input_label.setText(f"Eingabeordner: {folder}")

    def select_output_folder(self):
        folder = QtWidgets.QFileDialog.getExistingDirectory(self, "Ausgabeordner auswählen")
        if folder:
            self.output_folder = folder
            self.output_label.setText(f"Ausgabeordner: {folder}")

    def select_image(self):
        file, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Stempelbild auswählen", "",
                                                        "Bilder (*.png *.jpg *.jpeg)")
        if file:
            self.image_stamp_path = file
            self.image_label.setText(f"Stempelbild: {file}")

    def process_files(self):
        if not self.input_folder or not self.output_folder or not self.image_stamp_path:
            self.status_bar.showMessage("Bitte alle Verzeichnisse und das Stempelbild auswählen.", 5000)
            return

        self.status_bar.showMessage("Verarbeite Dateien...")

        # Finde alle .docx-Dateien im Eingabeordner
        docx_files = glob.glob(os.path.join(self.input_folder, "*.docx"))

        if not docx_files:
            self.status_bar.showMessage("Keine .docx-Dateien im Eingabeordner gefunden.", 5000)
            return

        for docx_file in docx_files:
            filename = os.path.basename(docx_file)
            output_docx_path = os.path.join(self.output_folder, filename)
            output_pdf_path = os.path.join(self.output_folder, os.path.splitext(filename)[0] + ".pdf")

            try:
                self.add_stamp_and_signature(docx_file, self.image_stamp_path, output_docx_path, output_pdf_path)
                self.status_bar.showMessage(f"{filename} erfolgreich verarbeitet.", 5000)
            except Exception as e:
                self.status_bar.showMessage(f"Fehler bei {filename}: {str(e)}", 5000)

    def add_stamp_and_signature(self, docx_path, image_path, output_docx_path, output_pdf_path):
        # Öffne das bestehende .docx-Dokument
        doc = Document(docx_path)

        # Füge 3 leere Zeilen am Ende des Dokuments ein
        for _ in range(3):
            doc.add_paragraph("")

        # Füge das Bild (Stempel mit Unterschrift) rechtsbündig ein
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.5))  # Passe die Breite entsprechend an
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Rechtsbündige Ausrichtung

        # Speichere das modifizierte .docx-Dokument
        doc.save(output_docx_path)

        # Konvertiere das .docx in ein PDF mit LibreOffice (Windows-Version)
        libreoffice_path = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
        if not os.path.isfile(libreoffice_path):
            raise FileNotFoundError("LibreOffice wurde nicht gefunden. Bitte überprüfen Sie den Pfad zu soffice.exe.")

        subprocess.run([
            libreoffice_path, "--headless", "--convert-to", "pdf", output_docx_path, "--outdir",
            self.output_folder
        ], check=True)


# Hauptprogramm
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    main_window = DocxProcessorApp()
    main_window.show()
    sys.exit(app.exec_())
