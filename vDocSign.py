from docx import Document
from docx.shared import Inches
from PIL import Image
from docx2pdf import convert
import os


def add_stamp_and_signature(docx_path, image_path, output_docx_path, output_pdf_path):
    # Öffne das bestehende .docx-Dokument
    doc = Document(docx_path)

    # Füge eine neue Seite hinzu
    doc.add_page_break()

    # Füge das Bild (Stempel mit Unterschrift) am Ende ein
    doc.add_paragraph(" ")
    doc.add_picture(image_path, width=Inches(2.5))  # Breite anpassen

    # Speichere das modifizierte .docx-Dokument
    doc.save(output_docx_path)

    # Konvertiere das .docx in ein PDF
    convert(output_docx_path, output_pdf_path)
    print(f"PDF gespeichert als {output_pdf_path}")


# Pfade anpassen
docx_input_path = "input.docx"
image_stamp_path = "stempel.png"  # Pfad zum eingescannten Stempel mit Unterschrift
docx_output_path = "output_with_stamp.docx"
pdf_output_path = "output_with_stamp.pdf"


# Funktion ausführen
add_stamp_and_signature(docx_input_path, image_stamp_path, docx_output_path, pdf_output_path)
