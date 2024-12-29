from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
import os

def add_stamp_and_signature(docx_path, image_path, output_docx_path, output_pdf_path):
    # Öffne das bestehende .docx-Dokument
    doc = Document(docx_path)
    
    # Füge ein paar leere Zeilen ein, um den Stempel nach unten zu schieben (3 Zeilen)
    for _ in range(3):
        doc.add_paragraph("")  # Fügt eine leere Zeile ein
    
    # Füge das Bild ein und richte es rechtsbündig aus
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.5))  # Passe die Breite des Bildes an
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Rechtsbündige Ausrichtung

    # Speichere das modifizierte .docx-Dokument
    doc.save(output_docx_path)
    print(f"Modifiziertes DOCX gespeichert als {output_docx_path}")

    # Konvertiere das .docx in ein PDF mit LibreOffice
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", output_docx_path, "--outdir", os.path.dirname(output_pdf_path)],
            check=True
        )
        print(f"PDF gespeichert als {output_pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"Fehler bei der Konvertierung: {e}")

# Pfade anpassen
docx_input_path = "input.docx"
image_stamp_path = "stempel.png"  # Pfad zum eingescannten Stempel mit Unterschrift
docx_output_path = "output_with_stamp.docx"
pdf_output_path = "output_with_stamp.pdf"

# Funktion ausführen
add_stamp_and_signature(docx_input_path, image_stamp_path, docx_output_path, pdf_output_path)
