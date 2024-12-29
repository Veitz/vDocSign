import sys
import os
import subprocess
from PyQt5 import QtWidgets, QtGui
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocxProcessorApp(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Docx Stempel und Signatur")
        self.setGeometry(100, 100, 600, 400)

        # Initialisiere Pfade
        self.docx_input_path = ""
        self.image_stamp_path = ""
        self.docx_output_path = ""
        self.pdf_output_path = ""

        # Menüband erstellen
        self.create_menu()

        # Zentrales Widget
        self.central_widget = QtWidgets.QWidget()
        self.setCentralWidget(self.central_widget)

        # Layout
        layout = QtWidgets.QVBoxLayout()
        self.central_widget.setLayout(layout)

        # Labels und Buttons für Verzeichnisse
        self.docx_label = QtWidgets.QLabel("DOCX-Datei: Nicht ausgewählt")
        layout.addWidget(self.docx_label)

        self.select_docx_button = QtWidgets.QPushButton("DOCX-Datei auswählen")
        self.select_docx_button.clicked.connect(self.select_docx_file)
        layout.addWidget(self.select_docx_button)

        self.image_label = QtWidgets.QLabel("Stempelbild: Nicht ausgewählt")
        layout.addWidget(self.image_label)

        self.select_image_button = QtWidgets.QPushButton("Stempelbild auswählen")
        self.select_image_button.clicked.connect(self.select_image)
        layout.addWidget(self.select_image_button)

        # Button zum Starten der Verarbeitung
        self.process_button = QtWidgets.QPushButton("DOCX verarbeiten")
        self.process_button.clicked.connect(self.process_file)
        layout.addWidget(self.process_button)

        # Statusanzeige
        self.status_bar = QtWidgets.QStatusBar()
        self.setStatusBar(self.status_bar)

    def create_menu(self):
        menubar = self.menuBar()
        file_menu = menubar.addMenu("&Datei")

        exit_action = QtWidgets.QAction("Beenden", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

    def select_docx_file(self):
        file, _ = QtWidgets.QFileDialog.getOpenFileName(self, "DOCX-Datei auswählen", "", "DOCX Dateien (*.docx)")
        if file:
            self.docx_input_path = file
            self.docx_output_path = os.path.splitext(file)[0] + "_with_stamp.docx"
            self.pdf_output_path = os.path.splitext(file)[0] + "_with_stamp.pdf"
            self.docx_label.setText(f"DOCX-Datei: {file}")

    def select_image(self):
        file, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Stempelbild auswählen", "",
                                                        "Bilder (*.png *.jpg *.jpeg)")
        if file:
            self.image_stamp_path = file
            self.image_label.setText(f"Stempelbild: {file}")

    def process_file(self):
        if not self.docx_input_path or not self.image_stamp_path:
            self.status_bar.showMessage("Bitte DOCX-Datei und Stempelbild auswählen.", 5000)
            return

        try:
            self.add_stamp_and_signature(self.docx_input_path, self.image_stamp_path, self.docx_output_path,
                                         self.pdf_output_path)
            self.status_bar.showMessage(f"Erfolgreich verarbeitet: {self.pdf_output_path}", 5000)
        except Exception as e:
            self.status_bar.showMessage(f"Fehler: {str(e)}", 5000)

    def add_stamp_and_signature(self, docx_path, image_path, output_docx_path, output_pdf_path):
        # Öffne das bestehende .docx-Dokument
        doc = Document(docx_path)

        # Füge 3 leere Zeilen am Ende des Dokuments ein
        for _ in range(3):
            doc.add_paragraph("")

        # Füge das Bild (Stempel mit Unterschrift) rechtsbündig ein
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(2.5))  # Passe die Breite entsprechend an
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Rechtsbündige Ausrichtung

        # Speichere das modifizierte .docx-Dokument
        doc.save(output_docx_path)

        # Konvertiere das .docx in ein PDF mit LibreOffice (Windows-Version)
        libreoffice
