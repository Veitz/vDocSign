from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
import os
import glob


def add_stamp_and_signature(docx_path, image_path, output_docx_path, output_pdf_path):
    # Öffne das bestehende .docx-Dokument
    doc = Document(docx_path)

    # Füge 3 leere Zeilen am Ende des Dokuments ein
    for _ in range(3):
        doc.add_paragraph("")

    # Füge das Bild (Stempel mit Unterschrift) rechtsbündig ein
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(1.5))  # Passe die Breite entsprechend an
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Rechtsbündige Ausrichtung

    # Speichere das modifizierte .docx-Dokument
    doc.save(output_docx_path)
    print(f"Modifiziertes DOCX gespeichert als {output_docx_path}")

    # Konvertiere das .docx in ein PDF mit LibreOffice
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", output_docx_path, "--outdir",
             os.path.dirname(output_pdf_path)],
            check=True
        )
        print(f"PDF gespeichert als {output_pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"Fehler bei der Konvertierung: {e}")


def process_all_docx(input_folder, output_folder, image_stamp_path):
    # Erstelle den Ausgabeordner, falls er nicht existiert
    os.makedirs(output_folder, exist_ok=True)

    # Finde alle .docx-Dateien im Eingabeordner
    docx_files = glob.glob(os.path.join(input_folder, "*.docx"))

    if not docx_files:
        print("Keine .docx-Dateien im Eingabeordner gefunden.")
        return

    for docx_file in docx_files:
        # Bestimme den Dateinamen ohne Pfad
        filename = os.path.basename(docx_file)

        # Erstelle Pfade für die Ausgabedateien
        output_docx_path = os.path.join(output_folder, filename)
        output_pdf_path = os.path.join(output_folder, os.path.splitext(filename)[0] + ".pdf")

        # Füge den Stempel hinzu und konvertiere in PDF
        add_stamp_and_signature(docx_file, image_stamp_path, output_docx_path, output_pdf_path)


# Pfade anpassen
input_folder = "input-docxs"
output_folder = "output-docxs"
image_stamp_path = "stempel.png"  # Pfad zum eingescannten Stempel mit Unterschrift

# Funktion ausführen
process_all_docx(input_folder, output_folder, image_stamp_path)
